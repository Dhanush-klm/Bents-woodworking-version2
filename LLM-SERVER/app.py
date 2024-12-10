import os
import uuid
import re
import logging
from flask import Flask, render_template, request, jsonify, session
from werkzeug.utils import secure_filename
from docx import Document
from dotenv import load_dotenv
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain.schema import Document as LangchainDocument, BaseRetriever
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate
import langsmith
from flask_cors import CORS
import psycopg2
from psycopg2.extras import RealDictCursor
import base64
from tenacity import retry, stop_after_attempt, wait_fixed, retry_if_exception_type
import time
import json
import numpy as np
from typing import List
from pydantic import BaseModel, Field

class LLMResponseError(Exception):
    pass

class LLMResponseCutOff(LLMResponseError):
    pass

class LLMNoResponseError(LLMResponseError):
    pass

load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": ["http://localhost:5173", "http://localhost:5002","https://bents-frontend-server.vercel.app","https://bents-backend-server.vercel.app"]}})



# System instructions
SYSTEM_INSTRUCTIONS = """You are an AI assistant representing Jason Bent's woodworking expertise. Your role is to:
1. Analyze woodworking documents and provide clear, natural responses that sound like Jason Bent is explaining the concepts.
2. Convert technical content into conversational, easy-to-understand explanations.
3. Focus on explaining the core concepts and techniques rather than quoting directly from transcripts.
4. Always maintain a friendly, professional tone as if Jason Bent is speaking directly to the user.
5. For each key point or technique mentioned that has a corresponding video source, include all three:
   - Timestamp in the format {{timestamp:HH:MM:SS}}
   - Video title in the format {{title:EXACT Video Title}}
   - URL in the format {{url:EXACT YouTube URL}}
6. If there is no relevant video content for a specific topic, provide the explanation without including timestamp/title/URL markers.
7. Organize multi-part responses clearly with natural transitions.
8. Keep responses concise and focused on the specific question asked.
9. If information isn't available in the provided context, clearly state that.
10. Always respond in English, regardless of the input language.
11. Avoid using phrases like "in the video" or "the transcript shows" - instead, speak directly about the techniques and concepts.


Example format for source citations (only when video content exists):
{{timestamp:05:30}}{{title:Workshop Tour 2024}}{{url:https://youtube.com/watch?v=abc123}}
12.Response Structure and Formatting:
   - Use markdown formatting with clear hierarchical structure
   - Each major section must start with '### ' followed by a number and bold title
   - Format section headers as: ### 1. **Title Here**
   - Use bullet points (-) for detailed explanations under each section
   - Each bullet point must contain 2-3 sentences minimum with examples
   - Add blank lines between major sections only
   - Indent bullet points with proper spacing
   - Do NOT use bold formatting (**) or line breaks within bullet point content
   - Bold formatting should ONLY be used in section headers
   - Keep all content within a bullet point on the same line
   - Any asterisks (*) in the content should be treated as literal characters, not formatting
13. Formatting Example:
### 1. **Main Point Title**
    - Detailed explanation that includes practical context and thorough reasoning.
      This should include specific examples or scenarios to illustrate the point.
      Additional details can be added to ensure comprehensive understanding.
### 2. **Second Point Title**
    - First detailed explanation with practical application and context. This should
      include why this point matters and how to implement it effectively. Include
      specific examples that demonstrate the concept in action.
    - Second explanation point with additional details and examples. Make sure to
      provide clear reasoning and practical applications. Include specific scenarios
      where this knowledge would be particularly useful.
Remember:
- You are speaking as Jason Bent's AI assistant and so if you are mentioning jason bent, you should use the word "Jason Bent" instead of "I" like "Jason Bent will suggest that you..."
- Focus on analyzing the transcripts and explaining the concepts naturally rather than quoting transcripts
- Only provide timestamp, video title, and URL when there is actual video content available
- NEVER create or invent video references, timestamps, or URLs
- ONLY use video references that are explicitly provided in the source documents
- If explaining a concept without a matching video, simply provide the explanation without any video markers
- Keep responses clear, practical, and focused on woodworking expertise
"""
app.secret_key = os.urandom(24)  # Set a secret key for sessions

# Access your API keys (set these in environment variables)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_API_KEY"] = os.getenv("LANGSMITH_API_KEY")
os.environ["LANGCHAIN_ENDPOINT"] = "https://api.smith.langchain.com"
os.environ["LANGCHAIN_PROJECT"] = "jason-json"

# Initialize Langchain components
embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, model="gpt-4o-2024-11-20", temperature=0)

logging.basicConfig(level=logging.DEBUG)

def rewrite_query(query, chat_history=None):
    """
    Rewrites the user query to be more specific and searchable using LLM.
    """
    try:
        # Create prompt that's compatible with the existing LLM setup
        rewrite_prompt = f"""You are bent's woodworks assistant so question will be related to wood shop. Rewrites user query to make them more specific and searchable, taking into account the chat history if provided. Only return the rewritten query without any explanations.

        Original query: {query}
        
        Chat history: {json.dumps(chat_history) if chat_history else '[]'}
        
        Rewritten query:"""
        
        # Use the existing LLM instance
        response = llm.predict(rewrite_prompt)
        
        # Clean up the response
        cleaned_response = response.replace("Rewritten query:", "").strip()
        
        # Add logging for debugging
        logging.debug(f"Original query: {query}")
        logging.debug(f"Rewritten query: {cleaned_response}")
        
        return cleaned_response if cleaned_response else query
        
    except Exception as e:
        logging.error(f"Error in query rewriting: {str(e)}", exc_info=True)
        return query  # Fallback to original query

def get_matched_products(video_title):
    logging.debug(f"Attempting to get matched products for title: {video_title}")
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            query = """
                SELECT id, title, tags, link FROM products 
                WHERE LOWER(tags) LIKE LOWER(%s)
            """
            search_term = f"%{video_title}%"
            logging.debug(f"Executing SQL query: {query} with search term: {search_term}")
            cur.execute(query, (search_term,))
            matched_products = cur.fetchall()
            logging.debug(f"Raw matched products from database: {matched_products}")
        conn.close()

        related_products = [
            {
                'id': product['id'],
                'title': product['title'],
                'tags': product['tags'].split(',') if product['tags'] else [],
                'link': product['link']
            } for product in matched_products
        ]

        logging.debug(f"Processed related products: {related_products}")
        return related_products

    except Exception as e:
        logging.error(f"Error in get_matched_products: {str(e)}", exc_info=True)
        return []
def verify_database():
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT COUNT(*) FROM products")
            count = cur.fetchone()['count']
            logging.info(f"Total products in database: {count}")
            
            cur.execute("SELECT title FROM products LIMIT 5")
            sample_titles = [row['title'] for row in cur.fetchall()]
            logging.info(f"Sample product titles: {sample_titles}")
        conn.close()
        return True
    except Exception as e:
        logging.error(f"Database verification failed: {str(e)}", exc_info=True)
        return False

def process_answer(answer, urls, source_documents):
    def extract_context(text, marker_pos, window=150):
        start = max(0, marker_pos - window)
        end = min(len(text), marker_pos + window)
        return text[start:end].strip()

    def generate_description(context, timestamp):
        description_prompt = f"""
        Given this woodworking video context at {timestamp}, create an extremely concise action phrase (max 6-8 words).

        Context: {context}

        Rules:
        1. Start with an action verb
        2. Name the specific tool/technique
        3. Must be 6-8 words only
        4. Focus on the single main action
        5. Be direct and clear

        Example formats:
        - "Demonstrates table saw fence alignment technique"
        - "Installs dust collection system components"
        - "Shows track saw cutting method"

        Description:"""

        try:
            enhanced_description = llm.predict(description_prompt).strip()
            words = enhanced_description.split()[:8]
            return ' '.join(words)
        except Exception as e:
            logging.error(f"Error generating description: {str(e)}")
            return ' '.join(context.split()[:6])

    def process_markers(timestamp_match, title_match, source_index):
        timestamp = timestamp_match.group(1)
        title = title_match.group(1)
        url_match = re.search(r'\{url:([^\}]+)\}', answer[timestamp_match.start():])
        url = url_match.group(1) if url_match else None
        
        timestamp_pos = timestamp_match.start()
        context = extract_context(answer, timestamp_pos)
        enhanced_description = generate_description(context, timestamp)
        
        full_urls = [combine_url_and_timestamp(url, timestamp)] if url else []
        
        return {
            'links': full_urls,
            'timestamp': timestamp,
            'description': enhanced_description,
            'video_title': title
        }
    
    # Find all timestamp and title markers
    timestamp_matches = list(re.finditer(r'\{timestamp:([^\}]+)\}', answer))
    title_matches = list(re.finditer(r'\{title:([^\}]+)\}', answer))
    
    # Ensure we have matching pairs
    timestamps_info = []
    for i, (ts_match, title_match) in enumerate(zip(timestamp_matches, title_matches)):
        timestamps_info.append(process_markers(ts_match, title_match, i))
    
    video_dict = {
        str(i): {
            'urls': entry['links'],
            'timestamp': entry['timestamp'],
            'description': entry['description'],
            'video_title': entry['video_title']
        }
        for i, entry in enumerate(timestamps_info)
    }
    
    # Clean up the answer text
    processed_answer = re.sub(r'\{timestamp:[^\}]+\}', '', answer)
    processed_answer = re.sub(r'\{title:[^\}]+\}', '', processed_answer)
    processed_answer = re.sub(r'\{url:[^\}]+\}', '', processed_answer)
    processed_answer = re.sub(r'\[?video\s*\d+\]?', '', processed_answer, flags=re.IGNORECASE)
    
    return processed_answer, video_dict

def combine_url_and_timestamp(base_url, timestamp):
    parts = timestamp.split(':')
    if len(parts) == 2:
        minutes, seconds = map(int, parts)
        total_seconds = minutes * 60 + seconds
    elif len(parts) == 3:
        hours, minutes, seconds = map(int, parts)
        total_seconds = hours * 3600 + minutes * 60 + seconds
    else:
        raise ValueError("Invalid timestamp format")

    if '?' in base_url:
        return f"{base_url}&t={total_seconds}"
    else:
        return f"{base_url}?t={total_seconds}"

def extract_text_from_docx(file):
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_metadata_from_text(text):
    title = text.split('\n')[0] if text else "Untitled Video"
    return {"title": title}

def upsert_transcript(transcript_text, metadata, index_name):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_text(transcript_text)
    
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor() as cur:
            for i, chunk in enumerate(chunks):
                chunk_metadata = metadata.copy()
                chunk_metadata['chunk_id'] = f"{metadata['title']}_chunk_{i}"
                chunk_metadata['url'] = metadata.get('url', '')
                chunk_metadata['title'] = metadata.get('title', 'Unknown Video')
                
                # Generate embeddings for the chunk
                chunk_embedding = embeddings.embed_query(chunk)
                
                # Insert into bents table
                cur.execute("""
                    INSERT INTO bents (text, title, url, chunk_id, vector)
                    VALUES (%s, %s, %s, %s, %s)
                    ON CONFLICT (chunk_id) DO UPDATE
                    SET text = EXCLUDED.text, vector = EXCLUDED.vector
                """, (chunk, chunk_metadata['title'], chunk_metadata['url'], 
                      chunk_metadata['chunk_id'], str(chunk_embedding)))
        conn.commit()
    except Exception as e:
        logging.error(f"Error upserting transcript: {str(e)}")
        raise
    finally:
        if conn:
            conn.close()

@retry(stop=stop_after_attempt(3), wait=wait_fixed(2), retry=retry_if_exception_type(LLMResponseError))
def retry_llm_call(qa_chain, query, chat_history):
    try:
        result = qa_chain({"question": query, "chat_history": chat_history})
        
        if result is None or 'answer' not in result or not result['answer']:
            raise LLMNoResponseError("LLM failed to generate a response")
        
        if result['answer'].endswith('...') or len(result['answer']) < 20:
            raise LLMResponseCutOff("LLM response appears to be cut off")
        return result
    except Exception as e:
        if isinstance(e, LLMResponseError):
            logging.error(f"LLM call failed: {str(e)}")
            raise
        logging.error(f"Unexpected error in LLM call: {str(e)}")
        raise LLMNoResponseError("LLM failed due to an unexpected error")

def connect_to_db():
    return psycopg2.connect(os.getenv("POSTGRES_URL"))

def get_embeddings(query):
    try:
        return embeddings.embed_query(query)
    except Exception as e:
        logging.error(f"Error generating embeddings: {str(e)}")
        raise

def cosine_similarity(v1, v2):
    return np.dot(v1, v2) / (np.linalg.norm(v1) * np.linalg.norm(v2))

def search_neon_db(query_embedding, table_name, top_k=5):
    conn = None
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            # Query with table name as parameter
            query = f"""
                SELECT id, vector, text, title, url, chunk_id 
                FROM {table_name}
                WHERE vector IS NOT NULL
            """
            cur.execute(query)
            rows = cur.fetchall()
            
            similarities = []
            for row in rows:
                try:
                    vector_str = row['vector']
                    vector_values = vector_str.strip('[]').split(',')
                    vector = np.array([float(x.strip()) for x in vector_values])
                    
                    if len(vector) == len(query_embedding):
                        similarity = cosine_similarity(query_embedding, vector)
                        similarities.append((similarity, row))
                except Exception as e:
                    logging.error(f"Error processing vector for row {row['id']}: {str(e)}")
                    continue
            
            similarities.sort(reverse=True, key=lambda x: x[0])
            return [
                {
                    'id': row['id'],
                    'text': row['text'],
                    'title': row['title'],
                    'url': row['url'],
                    'chunk_id': row['chunk_id'],
                    'similarity_score': float(sim)
                }
                for sim, row in similarities[:top_k]
            ]

    except Exception as e:
        logging.error(f"Error in search_neon_db: {str(e)}")
        raise
    finally:
        if conn:
            conn.close()

def handle_query(query):
    query_embedding = get_embeddings(query)
    results = search_neon_db(query_embedding)
    return results

# Update the custom retriever class
class CustomNeonRetriever(BaseRetriever, BaseModel):
    table_name: str = Field(...)  # The ... means this field is required
    
    class Config:
        arbitrary_types_allowed = True  # This allows for non-pydantic types
    
    def get_relevant_documents(self, query: str) -> List[LangchainDocument]:
        query_embedding = get_embeddings(query)
        results = search_neon_db(query_embedding, self.table_name)
        
        documents = []
        for result in results:
            timestamp_match = re.search(r'\[Timestamp: ([^\]]+)\]', result['text'])
            timestamp = timestamp_match.group(1) if timestamp_match else None
            
            doc = LangchainDocument(
                page_content=result['text'],
                metadata={
                    'title': result['title'],
                    'url': result['url'],
                    'timestamp': timestamp,
                    'chunk_id': result['chunk_id'],
                    'source': self.table_name
                }
            )
            documents.append(doc)
        
        return documents

    async def aget_relevant_documents(self, query: str) -> List[LangchainDocument]:
        return self.get_relevant_documents(query)

def get_all_related_products(video_dict):
    """Get related products from all video titles in video_links"""
    all_products = []  # Use list instead of set
    seen_products = set()  # Use a set of IDs to track duplicates
    
    # Extract unique video titles from video_dict
    video_titles = {entry['video_title'] for entry in video_dict.values()}
    
    # Get products for each video title
    for title in video_titles:
        if title:
            products = get_matched_products(title)
            for product in products:
                # Use product ID as unique identifier
                if product['id'] not in seen_products:
                    seen_products.add(product['id'])
                    all_products.append(product)
    
    return all_products

@app.route('/')
@app.route('/database')
def serve_spa():
    return render_template('index.html')

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.json
        user_query = data['message'].strip()
        chat_history = data.get('chat_history', [])

        # Format chat history
        formatted_history = []
        for i in range(0, len(chat_history) - 1, 2):
            human = chat_history[i]
            ai = chat_history[i + 1] if i + 1 < len(chat_history) else ""
            formatted_history.append((human, ai))

        # Add relevancy check
        relevance_check_prompt = f"""
        Given the following question or message and the chat history, determine if it is:
        1. A greeting or send-off like "thankyou" or "goodbye" or messages or casual messages like 'hey' or 'hello' or general conversation starter
        2. Related to woodworking, tools, home improvement, or the assistant's capabilities and also query about bents-woodworking youtube channel general questions.
        3. Related to the company, its products, services, or business operations
        4. A continuation or follow-up question to the previous conversation
        5. Related to violence, harmful activities, or other inappropriate content
        6. Completely unrelated to the above topics and not a continuation of the conversation
        7. if user is asking about jason bents.

        If it falls under category 1, respond with 'GREETING'.
        If it falls under categories 2, 3, 4 or 7 respond with 'RELEVANT'.
        If it falls under category 5, respond with 'INAPPROPRIATE'.
        If it falls under category 6, respond with 'NOT RELEVANT'.

        Chat History:
        {formatted_history[-3:] if formatted_history else "No previous context"}

        Current Question: {user_query}
        
        Response (GREETING, RELEVANT, INAPPROPRIATE, or NOT RELEVANT):
        """
        
        relevance_response = llm.predict(relevance_check_prompt)
        
        # Handle non-relevant cases using LLM
        if "GREETING" in relevance_response.upper():
            greeting_prompt = f"""
            The following message is a greeting or casual message. Please provide a friendly and engaging response.

            Message: {user_query}

            Response:
            """
            greeting_response = llm.predict(greeting_prompt)
            return jsonify({
                'response': greeting_response,
                'related_products': [],
                'urls': [],
                'contexts': [],
                'video_links': {}
            })
        elif "INAPPROPRIATE" in relevance_response.upper():
            inappropriate_prompt = f"""
            The following message is inappropriate or related to harmful activities. Please provide a polite and firm response indicating the limitations of the assistant.

            Message: {user_query}

            Response:
            """
            inappropriate_response = llm.predict(inappropriate_prompt)
            return jsonify({
                'response': inappropriate_response,
                'related_products': [],
                'urls': [],
                'contexts': [],
                'video_links': {}
            })
        elif "NOT RELEVANT" in relevance_response.upper():
            not_relevant_prompt = f"""
             The following question is not directly related to woodworking or the assistant's expertise. Provide a direct response that:
             1. Politely acknowledges the question
             2. Explains that you are specialized in woodworking and Jason Bent's content
             3. Asks them to rephrase their question to relate to woodworking topics
    
             Question: {user_query}

             Response (start directly with your message):
             """
             not_relevant_response = llm.predict(not_relevant_prompt)
             return jsonify({
                'response': not_relevant_response.strip(),
                'related_products': [],
                'urls': [],
                'contexts': [],
                'video_links': {}
    })

        # Only proceed with query rewriting if the query is relevant
        rewritten_query = rewrite_query(user_query, formatted_history)
        logging.debug(f"Query rewritten from '{user_query}' to '{rewritten_query}'")

        # Continue with your existing retrieval and response generation logic...
        retriever = CustomNeonRetriever(table_name="bents")
        
        # Define prompt
        prompt = ChatPromptTemplate.from_messages([
            SystemMessagePromptTemplate.from_template(SYSTEM_INSTRUCTIONS),
            HumanMessagePromptTemplate.from_template(
                "Context: {context}\n\nChat History: {chat_history}\n\nQuestion: {question}\n\n"
                "Instruction: Only use the provided context to generate the answer."
            )
        ])
        
        qa_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=retriever,
            combine_docs_chain_kwargs={"prompt": prompt},
            return_source_documents=True
        )
        
        # Use the rewritten query instead of the original
        result = qa_chain({"question": rewritten_query, "chat_history": formatted_history})
        
        # Extract answer and source documents
        raw_answer = result['answer']  # Store the raw answer before processing
        source_documents = result['source_documents']
        
        # Process source documents
        urls = []
        contexts = []
        for doc in source_documents:
            if 'url' in doc.metadata:
                urls.append(doc.metadata['url'])
            contexts.append(doc.page_content)
        
        # Process the answer to get video dictionary
        processed_answer, video_dict = process_answer(raw_answer, urls, source_documents)
        
        response_data = {
            'response': processed_answer,
            'raw_response': raw_answer,  # Include the raw LLM response
            'video_links': video_dict,
            'related_products': get_all_related_products(video_dict),
            'urls': urls
        }
        
        # Only include 'Related Videos' if there are actual video links
        if video_dict:
            response_data['related_videos'] = video_dict
        
        return jsonify(response_data)
        
    except Exception as e:
        logging.error(f"Error in chat route: {str(e)}", exc_info=True)
        return jsonify({'error': 'An error occurred processing your request'}), 500

@app.route('/api/user/<user_id>', methods=['GET'])
def get_user_data(user_id):
    try:
        user_data = {
            'conversationsBySection': {
                "bents": []
            },
            'searchHistory': [],
            'selectedIndex': "bents"
        }
        return jsonify(user_data)
    except Exception as e:
        logging.error(f"Error fetching user data: {str(e)}", exc_info=True)
        return jsonify({'error': 'An error occurred fetching user data'}), 500

@app.route('/documents')
def get_documents():
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT * FROM products")
            documents = cur.fetchall()
        conn.close()
        return jsonify(documents)
    except Exception as e:
        print(f"Error in get_documents: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/add_document', methods=['POST'])
def add_document():
    data = request.json
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute(
                "INSERT INTO products (title, tags, link) VALUES (%s, %s, %s) RETURNING id",
                (data['title'], ','.join(data['tags']), data['link'])
            )
            product_id = cur.fetchone()['id']
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'product_id': product_id})
    except Exception as e:
        print(f"Error in add_document: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/delete_document', methods=['POST'])
def delete_document():
    data = request.json
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor() as cur:
            cur.execute("DELETE FROM products WHERE id = %s", (data['id'],))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        print(f"Error in delete_document: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/update_document', methods=['POST'])
def update_document():
    data = request.json
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE products SET title = %s, tags = %s, link = %s WHERE id = %s",
                (data['title'], ','.join(data['tags']), data['link'], data['id'])
            )
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        print(f"Error in update_document: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/search', methods=['POST'])
def search():
    try:
        data = request.json
        query = data.get('query', '').strip()

        if not query:
            return jsonify({'error': 'Query is required'}), 400

        # Generate embeddings for the query
        query_embedding = get_embeddings(query)
        
        # Get raw results from database
        results = search_neon_db(query_embedding)
        
        # Return only the database results
        return jsonify({
            'results': results,
            'count': len(results)
        }), 200

    except Exception as e:
        logging.error(f"Error in search route: {str(e)}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/upload_document', methods=['POST'])
def upload_document():
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file part'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No selected file'})
    
    table_name = "bents"
    
    if file and file.filename.endswith('.docx'):
        try:
            filename = secure_filename(file.filename)
            file_path = os.path.join('/tmp', filename)
            file.save(file_path)
            
            transcript_text = extract_text_from_docx(file_path)
            metadata = extract_metadata_from_text(transcript_text)
            
            # Generate embeddings for the text
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            chunks = text_splitter.split_text(transcript_text)
            
            conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
            with conn.cursor() as cur:
                for i, chunk in enumerate(chunks):
                    chunk_embedding = embeddings.embed_query(chunk)
                    chunk_id = f"{metadata['title']}_chunk_{i}"
                    
                    cur.execute(f"""
                        INSERT INTO {table_name} (text, title, url, chunk_id, vector)
                        VALUES (%s, %s, %s, %s, %s)
                        ON CONFLICT (chunk_id) DO UPDATE
                        SET text = EXCLUDED.text, vector = EXCLUDED.vector
                    """, (
                        chunk,
                        metadata.get('title', 'Unknown Video'),
                        metadata.get('url', ''),
                        chunk_id,
                        str(chunk_embedding)
                    ))
            
            conn.commit()
            conn.close()
            os.remove(file_path)
            
            return jsonify({'success': True, 'message': 'File uploaded and processed successfully'})
            
        except Exception as e:
            logging.error(f"Error processing document: {str(e)}")
            return jsonify({'success': False, 'message': f'Error processing document: {str(e)}'})
    else:
        return jsonify({'success': False, 'message': 'Invalid file format'})

if __name__ == '__main__':
    verify_database()
    app.run(debug=True, port=5000)
